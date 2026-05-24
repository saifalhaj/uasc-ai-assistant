import io

import docx
import pdfplumber

from .base import ParsedDocument, Source


class UploadSource(Source):
    """Handles user-uploaded PDF, DOCX, TXT files."""

    @property
    def name(self) -> str:
        return "upload"

    @property
    def classification(self) -> str:
        return "public"

    @property
    def tier(self) -> str:
        return "open"

    @property
    def mode(self) -> str:
        return "upload"

    async def parse(self, raw: bytes, content_type: str) -> ParsedDocument:
        if content_type == "application/pdf" or content_type.endswith(".pdf"):
            text = self._parse_pdf(raw)
        elif content_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            text = self._parse_docx(raw)
        else:
            text = raw.decode("utf-8", errors="replace")

        return ParsedDocument(text=text, metadata={"source": "upload"})

    def _parse_pdf(self, raw: bytes) -> str:
        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)
        return "\n\n".join(parts)

    def _parse_docx(self, raw: bytes) -> str:
        doc = docx.Document(io.BytesIO(raw))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
